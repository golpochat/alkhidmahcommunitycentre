"""Generate client content requirements DOCX for Al Khidmah Mosque project."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement


OUTPUT = Path(__file__).resolve().parents[1] / "docs" / "client-content-requirements.docx"

HEADERS = [
    "Section",
    "Field",
    "Priority",
    "Current Placeholder",
    "Your Text / Correct Details",
    "Where to Update",
]

ROWS = [
    # 1 Organisation
    ("1. Organisation Identity", "Official organisation / site name", "Required", "Al Khidmah Community Centre", "", "Super Admin → Settings → Site"),
    ("1. Organisation Identity", "Registered charity number", "Required", "CHY 22345", "", "Super Admin → Settings → Site"),
    ("1. Organisation Identity", "Live website URL (domain)", "Required", "https://alkhidmah.ie", "", "Super Admin → Settings → Site"),
    ("1. Organisation Identity", "Logo file (PNG/SVG, transparent background preferred)", "Required", "Placeholder logo in /public/logo/", "", "Super Admin → Settings → Site → Branding → Upload Logo"),
    ("1. Organisation Identity", "Favicon / app icon (square PNG or SVG)", "Recommended", "Gold minaret on emerald green background (/favicon.png)", "", "Super Admin → Settings → Site → Branding → Upload Favicon (or npm run favicon:generate)"),
    ("1. Organisation Identity", "SEO description (Google search snippet)", "Recommended", "Serving the Muslim community of Clondalkin with prayer, education, and community support.", "", "Ask developer (constants.ts / metadata)"),
    # 2 Contact
    ("2. Contact Details", "Full postal address", "Required", "Unit 4, Monastery Road, Clondalkin, Dublin 22, D22 YX82", "", "Super Admin → Settings → Site"),
    ("2. Contact Details", "Phone number", "Required", "+353 1 457 8900", "", "Super Admin → Settings → Site"),
    ("2. Contact Details", "Public email address", "Required", "info@alkhidmahmosque.ie", "", "Super Admin → Settings → Site"),
    ("2. Contact Details", "Staff notification email (donations, registrations, contact form)", "Required", "info@alkhidmahmosque.ie", "", "Super Admin → Settings → Site"),
    ("2. Contact Details", "WhatsApp number", "Recommended", "+353851234567", "", "Super Admin → Settings → Site"),
    ("2. Contact Details", "Facebook URL", "If applicable", "https://facebook.com/alkhidmahmosque", "", "Super Admin → Settings → Site"),
    ("2. Contact Details", "Instagram URL", "If applicable", "https://instagram.com/alkhidmahmosque", "", "Super Admin → Settings → Site"),
    ("2. Contact Details", "YouTube URL", "If applicable", "https://youtube.com/@alkhidmahmosque", "", "Super Admin → Settings → Site"),
    ("2. Contact Details", "X (Twitter) URL", "If applicable", "https://twitter.com/alkhidmahmosque", "", "Super Admin → Settings → Site"),
    ("2. Contact Details", "Opening hours — weekdays", "Required", "Monday – Friday: 9:00 AM – 9:00 PM", "", "Ask developer (contact.ts) or provide text"),
    ("2. Contact Details", "Opening hours — weekends", "Required", "Saturday – Sunday: 9:00 AM – 9:00 PM", "", "Ask developer (contact.ts) or provide text"),
    ("2. Contact Details", "Opening hours — prayer access note", "Recommended", "Prayer Times: Open for all daily prayers", "", "Ask developer (contact.ts) or provide text"),
    # 3 Home
    ("3. Home Page", "Main headline (H1)", "Required", "Serving the Muslim Community of Clondalkin", "", "Ask developer (hero-section.tsx)"),
    ("3. Home Page", "Hero tagline / sub-text", "Required", "A centre for worship, learning, and community — welcoming all to prayer, education, events, and charitable giving.", "", "Ask developer (hero-section.tsx)"),
    ("3. Home Page", "Hero background photo (mosque exterior/interior)", "Recommended", "Stock placeholder image", "", "Provide high-res photo to developer"),
    ("3. Home Page", "Explore section — heading", "Recommended", "More at the Centre", "", "Ask developer (home-explore-section.tsx)"),
    ("3. Home Page", "Explore section — intro line", "Recommended", "Donations, events, classes, and photos each have their own page.", "", "Ask developer (home-explore-section.tsx)"),
    ("3. Home Page", "Explore — Donate card description", "Recommended", "Support prayer, education, and community outreach.", "", "Ask developer (home-explore-section.tsx)"),
    ("3. Home Page", "Explore — Events card description", "Recommended", "Community gatherings, programmes, and special occasions.", "", "Ask developer (home-explore-section.tsx)"),
    ("3. Home Page", "Explore — Education card description", "Recommended", "Qur'an, Arabic, and Islamic classes for all ages.", "", "Ask developer (home-explore-section.tsx)"),
    ("3. Home Page", "Explore — Gallery card description", "Recommended", "Photos from prayer, education, and community life.", "", "Ask developer (home-explore-section.tsx)"),
    ("3. Home Page", "Footer & brand tagline", "Recommended", "Prayer, education, and community services for Clondalkin and surrounding areas.", "", "Ask developer (footer.tsx)"),
    ("3. Home Page", "About teaser — intro paragraph", "Recommended", "A cornerstone of the Muslim community in Clondalkin since 2010…", "", "Ask developer (about-teaser.tsx)"),
    ("3. Home Page", "About teaser — mission summary", "Recommended", "To serve the Muslim community of Clondalkin by providing a centre for the five daily prayers…", "", "Ask developer (about-teaser.tsx)"),
    # 4 About
    ("4. About Page", "Page intro (hero description)", "Required", "{Site name} has been a cornerstone of the Muslim community in Clondalkin since 2010…", "", "Ask developer (about page)"),
    ("4. About Page", "Our Mission — paragraph 1", "Required", "To serve the Muslim community of Clondalkin by providing a centre for the five daily prayers…", "", "Ask developer (about page)"),
    ("4. About Page", "Our Mission — paragraph 2", "Recommended", "We strive to nurture faith, knowledge, and compassion…", "", "Ask developer (about page)"),
    ("4. About Page", "Our History — paragraph 1", "Required", "Founded in 2010 by a group of dedicated community members…", "", "Ask developer (about page)"),
    ("4. About Page", "Our History — paragraph 2", "Required", "Today, we serve hundreds of families with daily prayers…", "", "Ask developer (about page)"),
    ("4. About Page", "Year founded", "Required", "2010", "", "Include in history text above"),
    ("4. About Page", "About page hero photo", "Recommended", "Stock placeholder", "", "Provide photo to developer"),
    ("4. About Page", "Community gathering photo (mission section)", "Recommended", "Stock placeholder", "", "Provide photo to developer"),
    # 5 Values
    ("5. Our Values", "Value 1 — title", "Required", "Education", "", "Admin → About Page → Our Values"),
    ("5. Our Values", "Value 1 — description", "Required", "Providing Quran and Islamic education for children and adults of all backgrounds.", "", "Admin → About Page → Our Values"),
    ("5. Our Values", "Value 2 — title", "Required", "Charity", "", "Admin → About Page → Our Values"),
    ("5. Our Values", "Value 2 — description", "Required", "Supporting the needy through zakah, sadaqah, and community welfare programmes.", "", "Admin → About Page → Our Values"),
    ("5. Our Values", "Value 3 — title", "Required", "Community", "", "Admin → About Page → Our Values"),
    ("5. Our Values", "Value 3 — description", "Required", "Building a welcoming space for worship, fellowship, and cultural connection.", "", "Admin → About Page → Our Values"),
    ("5. Our Values", "Value 4 — title", "Required", "Excellence", "", "Admin → About Page → Our Values"),
    ("5. Our Values", "Value 4 — description", "Required", "Upholding the highest standards as a registered and verified charity.", "", "Admin → About Page → Our Values"),
    # 6 Committee
    ("6. Mosque Committee", "Member 1 — full name", "Required", "Dr. Ibrahim Hassan", "", "Admin → About Page → Committee"),
    ("6. Mosque Committee", "Member 1 — role", "Required", "Chairperson", "", "Admin → About Page → Committee"),
    ("6. Mosque Committee", "Member 1 — short bio", "Recommended", "Leading the mosque committee with over 15 years of community service in Clondalkin.", "", "Admin → About Page → Committee"),
    ("6. Mosque Committee", "Member 1 — photo", "Optional", "(none)", "", "Admin → About Page → Committee"),
    ("6. Mosque Committee", "Member 2 — full name", "Required", "Fatima Al-Rashid", "", "Admin → About Page → Committee"),
    ("6. Mosque Committee", "Member 2 — role", "Required", "Secretary", "", "Admin → About Page → Committee"),
    ("6. Mosque Committee", "Member 2 — short bio", "Recommended", "Coordinating mosque operations and community communications…", "", "Admin → About Page → Committee"),
    ("6. Mosque Committee", "Member 3 — full name", "Required", "Yusuf Mahmoud", "", "Admin → About Page → Committee"),
    ("6. Mosque Committee", "Member 3 — role", "Required", "Treasurer", "", "Admin → About Page → Committee"),
    ("6. Mosque Committee", "Member 3 — short bio", "Recommended", "Managing charitable funds with transparency and accountability…", "", "Admin → About Page → Committee"),
    ("6. Mosque Committee", "Member 4 — full name", "Required", "Aisha O'Brien", "", "Admin → About Page → Committee"),
    ("6. Mosque Committee", "Member 4 — role", "Required", "Education Coordinator", "", "Admin → About Page → Committee"),
    ("6. Mosque Committee", "Member 4 — short bio", "Recommended", "Overseeing Quran classes and Islamic education programmes…", "", "Admin → About Page → Committee"),
    ("6. Mosque Committee", "Additional members", "As needed", "(Add/remove in admin as required)", "", "Admin → About Page → Committee"),
    # 7 Donations
    ("7. Donation Categories", "Sadaqah — title & description", "Required", "Voluntary charity that brings barakah and supports ongoing mosque services.", "", "Admin → Donations → Categories"),
    ("7. Donation Categories", "Fitrah — title & description", "Required", "Eid al-Fitr charity due before Eid prayer…", "", "Admin → Donations → Categories"),
    ("7. Donation Categories", "Mosque Development — title & description", "Required", "Help maintain and expand our facilities…", "", "Admin → Donations → Categories"),
    ("7. Donation Categories", "Ramadan — title & description", "Required", "Support iftar programmes, taraweeh, and Ramadan initiatives.", "", "Admin → Donations → Categories"),
    ("7. Donation Categories", "Dawah — title & description", "Required", "Fund outreach, literature, and educational programmes.", "", "Admin → Donations → Categories"),
    ("7. Donation Categories", "Zakah — title & description", "Required", "Fulfill your obligatory charity and support those in need.", "", "Admin → Donations → Categories"),
    ("7. Donation Categories", "Additional categories (if needed)", "As needed", "(Create in admin — starts unpublished)", "", "Admin → Donations → Categories"),
    ("7. Donation Categories", "Confirm which categories to publish at launch", "Required", "(All start unpublished after seed)", "", "Admin → Donations → Categories → Publish toggle"),
    # 8 Payments
    ("8. Payment Setup", "Stripe account — publishable key, secret key, webhook secret", "If using Stripe", "(Not configured)", "", "Super Admin → Settings → Payment → Stripe"),
    ("8. Payment Setup", "PayPal account — client ID, secret, live/sandbox mode", "If using PayPal", "(Not configured)", "", "Super Admin → Settings → Payment → PayPal"),
    ("8. Payment Setup", "Bank transfer — account name, bank name, IBAN, BIC, reference note", "If using bank transfer", "(Not configured)", "", "Super Admin → Settings → Payment → Bank Transfer"),
    ("8. Payment Setup", "Donation currency", "Required", "EUR", "", "Super Admin → Settings → Payment"),
    ("8. Payment Setup", "Allow donor to cover processing fees?", "Recommended", "Configurable per gateway", "", "Super Admin → Settings → Payment"),
    # 9 Email
    ("9. Email (SMTP)", "SMTP host, port, username, password", "Required", "(Not configured)", "", "Super Admin → Settings → Email"),
    ("9. Email (SMTP)", "From email address (must be authorised on SMTP account)", "Required", "(Not configured)", "", "Super Admin → Settings → Email"),
    ("9. Email (SMTP)", "Test email after setup", "Required", "Send test from admin to confirm delivery", "", "Super Admin → Settings → Email → Test Send"),
    # 10 Prayer
    ("10. Prayer Times & Location", "Mosque GPS coordinates (latitude / longitude)", "Required", "53.3217, -6.4064 (Clondalkin default)", "", "Ask developer (constants.ts) if location differs"),
    ("10. Prayer Times & Location", "Calculation method preference", "Recommended", "AlAdhan API default for Ireland", "", "Admin → Prayer Timetable"),
    ("10. Prayer Times & Location", "Jumu'ah prayer time(s)", "Required", "(Set in admin)", "", "Admin → Prayer Timetable → Jumu'ah"),
    ("10. Prayer Times & Location", "Eid prayer times (Fitr & Adha)", "Required", "(Set before each Eid)", "", "Admin → Prayer Timetable → Eid"),
    ("10. Prayer Times & Location", "Daily prayer overrides (if different from calculated times)", "As needed", "(Optional)", "", "Admin → Prayer Timetable"),
    ("10. Prayer Times & Location", "Ramadan timetable — start date, moon sighting notes, PDF", "Before Ramadan", "(Configure each year)", "", "Admin → Prayer Timetable → Ramadan"),
    ("10. Prayer Times & Location", "Monthly timetable PDF (optional homepage banner)", "As needed", "(Configure each month)", "", "Admin → Prayer Timetable → Monthly"),
    # 11 Eid page
    ("11. Eid Page", "Sunnah reminders text", "Recommended", "Perform ghusl, wear your best clothes, eat dates…", "", "Ask developer (eid page) or provide text"),
    ("11. Eid Page", "Parking guidance", "Recommended", "Please use nearby street parking respectfully…", "", "Ask developer (eid page) or provide text"),
    ("11. Eid Page", "Women's prayer area guidance", "Recommended", "A dedicated sisters' prayer area is available…", "", "Ask developer (eid page) or provide text"),
    ("11. Eid Page", "Takbeer reminder", "Recommended", "Recite the takbeer after Fajr on the day of Eid…", "", "Ask developer (eid page) or provide text"),
    # 12 Contact page
    ("12. Contact Page", "Welcome message", "Recommended", "We welcome your questions, feedback, and enquiries…", "", "Ask developer (contact page) or provide text"),
    ("12. Contact Page", "Contact page hero photo", "Optional", "Stock placeholder", "", "Provide photo to developer"),
    # 13 Staff
    ("13. Staff & Access", "Super admin email (primary account owner)", "Required", "(Set at deployment)", "", "Deployment env: ADMIN_EMAIL"),
    ("13. Staff & Access", "List of staff to invite (name, email, role)", "Required", "(Provide spreadsheet or list)", "", "Super Admin → Invitations"),
    ("13. Staff & Access", "Role assignments (Admin, Web Admin, Account Admin, Editor)", "Required", "(See role guide in FEATURES.md)", "", "Super Admin → Invitations / Users"),
    ("13. Staff & Access", "Who can manage legal policies", "Recommended", "Admin and Super Admin (legal.manage permission)", "", "Super Admin → Roles & Permissions"),
    # 14 Initial content
    ("14. Initial Content (Admin)", "Events — titles, dates, descriptions, photos", "Required", "Sample seed events (unpublished)", "", "Admin → Events → Publish when ready"),
    ("14. Initial Content (Admin)", "Education programmes — title, schedule, teacher, fees", "Required", "Sample seed programmes (unpublished)", "", "Admin → Education → Publish when ready"),
    ("14. Initial Content (Admin)", "Gallery — album names and real mosque photos", "Recommended", "Sample stock photos (unpublished)", "", "Admin → Gallery → Publish albums when ready"),
    ("14. Initial Content (Admin)", "TV display — welcome notice, rotating announcements", "Recommended", "(Not configured)", "", "Admin → Screen & Announcements"),
    ("14. Initial Content (Admin)", "TV display — PIN lock code (if used)", "Optional", "(Not configured)", "", "Admin → Screen & Announcements → Settings"),
    # 15 Assets
    ("15. Photography & Media", "Home hero image (landscape, min 1920×1080)", "Recommended", "Stock placeholder", "", "Send to developer or upload via gallery/events"),
    ("15. Photography & Media", "Page hero images (About, Events, Education, Gallery, Donations, Contact)", "Optional", "Stock placeholders", "", "Send to developer"),
    ("15. Photography & Media", "Event cover photos (one per event)", "Recommended", "Stock placeholders", "", "Admin → Events (when creating)"),
    ("15. Photography & Media", "Real gallery photos (community, Ramadan, Eid, classes, youth)", "Recommended", "Stock placeholders", "", "Admin → Gallery"),
    # 16 Legal policies
    ("16. Legal Policies", "Privacy Policy — full legal text", "Required", "GDPR draft template (unpublished)", "", "Admin → Legal Policies → Privacy Policy"),
    ("16. Legal Policies", "Cookie Policy — full legal text", "Required", "Cookie draft template (unpublished)", "", "Admin → Legal Policies → Cookie Policy"),
    ("16. Legal Policies", "Terms of Use — full legal text", "Required", "Terms draft template (unpublished)", "", "Admin → Legal Policies → Terms of Use"),
    ("16. Legal Policies", "Policy version numbers", "Recommended", "1.0 (draft)", "", "Admin → Legal Policies (per policy)"),
    ("16. Legal Policies", "Effective date & last reviewed date", "Recommended", "Auto-filled to today when editing; confirm before publish", "", "Admin → Legal Policies (per policy)"),
    ("16. Legal Policies", "Publish each policy after solicitor sign-off", "Required", "All start as unpublished drafts", "", "Admin → Legal Policies → Published toggle → Save"),
    ("16. Legal Policies", "Confirm consent checkboxes match published policies", "Required", "Contact, registration, and member register forms link to Privacy Policy; donations show terms notice; cookie banner links to Cookie Policy", "", "Review after publishing"),
    # 17 Domain
    ("17. Domain & Launch", "Domain name registered and DNS pointed to hosting", "Required", "(Client responsibility)", "", "Client / hosting provider"),
    ("17. Domain & Launch", "SSL certificate active on live domain", "Required", "(Automatic on Vercel)", "", "Hosting provider"),
    ("17. Domain & Launch", "Hard-refresh browser after favicon or branding changes", "Recommended", "Cache-busted via ?v=8 on default favicon", "", "Client browsers after go-live"),
]

NOT_IN_DOC = [
    "Events — create and publish ongoing in Admin → Events",
    "Education programmes — Admin → Education",
    "Gallery photos — Admin → Gallery",
    "Prayer times & timetables — Admin → Prayer Timetable",
    "TV display announcements — Admin → Screen & Announcements",
    "Donation transaction history — Admin → Donations → Transactions",
    "Contact form submissions — Admin → Contact Messages",
    "Class registrations — Admin → Registrations",
    "Member accounts — self-registration on public site",
    "UI labels, buttons, and login pages — standard, no client input needed",
]

LAUNCH_MINIMUM = [
    "Official site name, charity number, and live domain URL",
    "Correct address, phone, email, WhatsApp, and notification email",
    "Social media links (or confirm removal if not used)",
    "Opening hours",
    "Home page headline and tagline (or approve placeholder text)",
    "About page mission and history (your real story)",
    "Committee members — real names and roles",
    "Our Values — four cards reflecting your mosque",
    "Donation categories — descriptions and which to publish",
    "Payment gateway credentials (Stripe / PayPal / bank transfer as applicable)",
    "SMTP email configured and test email sent",
    "Jumu'ah and Eid prayer times confirmed",
    "Staff list for invitations",
    "Logo uploaded in Super Admin settings",
    "Favicon confirmed (default gold minaret on emerald green, or upload custom)",
    "Legal policies reviewed by solicitor and published (Privacy, Cookie, Terms)",
]


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc: Document, rows: list[tuple[str, ...]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(HEADERS))
    table.style = "Table Grid"
    table.autofit = False

    for idx, header in enumerate(HEADERS):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, "1F4E3D")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)

    for row_idx, row_data in enumerate(rows, start=1):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    widths = [Inches(1.05), Inches(1.55), Inches(0.75), Inches(1.45), Inches(1.45), Inches(1.25)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def build_document() -> Document:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    title = doc.add_heading("Al Khidmah Community Centre", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("Client Content Requirements")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(31, 78, 61)

    date_p = doc.add_paragraph(f"Generated: {date.today():%d %B %Y}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.runs[0].font.size = Pt(10)
    date_p.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    intro = doc.add_paragraph(
        "This document lists all content and information needed from the client to launch and "
        "operate the Al Khidmah website and admin platform. A Markdown version is also available "
        "at docs/client-content-requirements.md. Fill in the "
        '"Your Text / Correct Details" column and return this document, or update directly in '
        "the admin panel where indicated."
    )
    intro.runs[0].font.size = Pt(10)

    note = doc.add_paragraph()
    note.add_run("Priority key: ").bold = True
    note.add_run(
        "Required = must be confirmed before launch. "
        "Recommended = improves the site but can use placeholder text initially. "
        "If applicable = only needed if you use that feature."
    )
    for run in note.runs:
        run.font.size = Pt(10)

    doc.add_heading("Before Launch — Minimum Checklist", level=2)
    for item in LAUNCH_MINIMUM:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("What Is NOT in This Document", level=2)
    doc.add_paragraph(
        "The following are managed ongoing in the admin panel after launch — "
        "no need to fill in here upfront:"
    )
    for item in NOT_IN_DOC:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading(f"Content Checklist ({len(ROWS)} items)", level=2)
    add_table(doc, ROWS)

    doc.add_paragraph()
    footer = doc.add_paragraph(
        "After returning this document, the development team will update any developer-managed "
        "fields (home page copy, opening hours, hardcoded page text, coordinates). "
        "All admin-managed items can be entered directly by staff once accounts are created."
    )
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
